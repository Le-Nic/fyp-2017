import pandas as pd
from docx import Document
from docx.shared import RGBColor	# font color
from docx.shared import Pt	# font size
from docx.oxml.shared import OxmlElement, qn	# toc
from docx.enum.text import WD_ALIGN_PARAGRAPH	# right align
from docx.oxml.ns import nsdecls	# table shading
from docx.oxml import parse_xml	# raw xml editing
from src import warehouse


extSupport = ['docx']

def genDOCX(colr, fileName, merged, isPDF=False):

	document = Document()

	# Microsoft Word TOC obtained directly from GitHub (https://github.com/python-openxml/python-docx/issues/36) by user scanny.
	paragraph = document.add_paragraph()
	run = paragraph.add_run()
	fldChar = OxmlElement('w:fldChar')	# creates a new element
	fldChar.set(qn('w:fldCharType'), 'begin')	# sets attribute on element
	instrText = OxmlElement('w:instrText')
	instrText.set(qn('xml:space'), 'preserve')	# sets attribute on element
	instrText.text = r'TOC \o "1-3" \h \z \u'	# change 1-3 depending on heading levels

	fldChar2 = OxmlElement('w:fldChar')
	fldChar2.set(qn('w:fldCharType'), 'separate')
	fldChar3 = OxmlElement('w:t')
	fldChar3.text = "Right-click to update field."
	fldChar2.append(fldChar3)

	fldChar4 = OxmlElement('w:fldChar')
	fldChar4.set(qn('w:fldCharType'), 'end')

	r_element = run._r
	r_element.append(fldChar)
	r_element.append(instrText)
	r_element.append(fldChar2)
	r_element.append(fldChar4)
	p_element = paragraph._p

	paragraph_format = document.styles['Normal'].paragraph_format
	paragraph_format.space_after = Pt(0)

	if merged:
		# MERGED vuln.
		document.add_page_break()
		paragraph = document.add_heading()
		run = paragraph.add_run('Consolidated')
		run.font.size = Pt(24)

		for ip, clusters in merged.items():

			if clusters:
				# IP Address
				paragraph = document.add_heading(level=2)
				run = paragraph.add_run('\n' + ip)
				run.font.size = Pt(20)

				for num, cluster in clusters.items():

					# Index
					paragraph = document.add_heading(level=3)
					run = paragraph.add_run(str(num +1))
					run.font.size = Pt(16)

					cveCluster = {}	# {cve: num of associated scanners}
					refCluster = {}	# {cve: num of associated scanners}

					scannerNames = []	# all scanners in a cluster
					for df in cluster.keys():
						if df.scanner not in scannerNames:
							scannerNames.append(df.scanner)

					# ----- Vulnerabilities -----
					for df, indexes in cluster.items():
						scannerColor = RGBColor(colr[df.scanner][0], colr[df.scanner][1], colr[df.scanner][2])

						for index, cves in indexes.items():

							# ----- Name -----
							if df.data['name'][index]:
								paragraph = document.add_paragraph()
								paragraph.style = document.styles['Title']
								run = paragraph.add_run(df.data['name'][index])
								run.font.color.rgb = scannerColor
								run.font.size = Pt(16)

							# CVEs
							for cve in cves:

								if cve not in cveCluster:
									cveCluster.update({
										cve: {
											scannerName: False for scannerName in scannerNames
										}
									})
								cveCluster[cve][df.scanner] = True

							# Refs
							if df.data['ref'][index]:
								for ref in df.data['ref'][index]:
									refType, refID = ref.split(':', 1)

									if refType != 'CVE':

										if refType not in refCluster:
											refCluster.update({
												refType: {
													refID: {
														scannerName: False for scannerName in scannerNames
													}
												}
											})

										else:
											if refID not in refCluster[refType]:
												refCluster[refType].update({
													refID: {
														scannerName: False for scannerName in scannerNames
													}
												})
										refCluster[refType][refID][df.scanner] = True

					paragraph = document.add_heading()
					run = paragraph.add_run('References')
					run.font.size = Pt(16)

					table = document.add_table(rows = len(cveCluster), cols = len(scannerNames) +1)
					table.autofit = True
					# ----- CVEs -----
					for cve, row in zip(sorted(cveCluster.keys()), table.rows):
						row.cells[0].text = cve

						for scan, cell in zip(cveCluster[cve], row.cells[1:]):

							if cveCluster[cve][scan]:
								cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
									nsdecls('w'),
									str(hex(colr[scan][0]))[2:] + str(hex(colr[scan][1]))[2:] + str(hex(colr[scan][2]))[2:]
								)))

					# ----- Refs -----
					for refType, refIDs in refCluster.items():
						paragraph = document.add_paragraph()
						run = paragraph.add_run('\n'+refType)
						run.bold = True

						table = document.add_table(rows = 1, cols = len(scannerNames) +1)
						table.autofit = True
						for refID, row in zip(refIDs.keys(), table.rows):

							row.cells[0].text = refID

							for scan, cell in zip(refCluster[refType][refID], row.cells[1:]):

								if refCluster[refType][refID][scan]:
									cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
										nsdecls('w'),
										str(hex(colr[scan][0]))[2:] + str(hex(colr[scan][1]))[2:] + str(hex(colr[scan][2]))[2:]
									)))

	for vuln in warehouse.vulnList:

		document.add_page_break()

		# set color for different scanner
		scannerColor = RGBColor(colr[vuln.scanner][0], colr[vuln.scanner][1], colr[vuln.scanner][2])

		# ----- IP Address -----
		paragraph = document.add_heading()
		run = paragraph.add_run(vuln.ip)
		run.font.color.rgb = scannerColor
		run.font.size = Pt(24)

		# ----- Host Name -----
		if vuln.host is not None:	# verbose off

			paragraph = document.add_paragraph()
			run = paragraph.add_run('Hostname\t\t')
			run.font.color.rgb = scannerColor
			paragraph.add_run(vuln.host)

		# ----- Operating System -----
		if vuln.os is not None:	# verbose off

			paragraph = document.add_paragraph()
			run = paragraph.add_run('Operating System\t')
			run.font.color.rgb = scannerColor

			if isinstance(vuln.os, list) is True:	# list
				paragraph.add_run('\n\t\t\t'.join(vuln.os))
			else:
				paragraph.add_run(vuln.os)

		# ----- Vulnerabilities -----
		for index, row in vuln.data.iterrows():
			# skip vulnerability if merged
			skip = False
			'''
			if merged:
				tempIP = merged.get(vuln.ip)
				if tempIP is not None:
					for cluster in tempIP:
						tempVuln = tempIP[cluster].get(vuln)
						if tempVuln is not None:
							tempIndex = tempVuln.get(index)
							if tempIndex is not None:
								skip = True
			'''
			if not skip:
				if row['name'] is not None and row['name'] != '':
					# ----- Name, Severity & Risk -----
					paragraph = document.add_heading(level=2)
					run = paragraph.add_run('\n'+row['name'])
					run.font.color.rgb = scannerColor
					run.font.size = Pt(16)

					if row['severity']:
						run = paragraph.add_run('\tSeverity: '+row['severity'])
						run.font.color.rgb = RGBColor(0xc0, 0x00, 0x00)
						run.font.size = Pt(11)

						if 'risk' in vuln.data.columns:
							if row['risk'] is not None:
								run = paragraph.add_run(' ({})'.format(row['risk']))
								run.font.color.rgb = RGBColor(0xc0, 0x00, 0x00)
								run.font.size = Pt(11)

					# ----- Port -----
					paragraph = document.add_paragraph()
					paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

					ports = row['port']
					first = True

					if ports:
						for port in ports:

							if not first:
								run = paragraph.add_run(' | ')
								run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
							else:
								first = False

							run = paragraph.add_run(port)
							run.font.color.rgb = scannerColor
							run.font.size = Pt(11)

					# ----- Summary, Desription & Solution -----
					groupColumns = ['summary', 'description', 'solution']
					first = True
					for column in groupColumns:

						if column in vuln.data.columns:
							paragraph = document.add_paragraph()

							if row[column]:
								if first:
									first = False
									run = paragraph.add_run(column.title())
								else:
									run = paragraph.add_run('\n' + column.title())

								run.font.color.rgb = scannerColor
								run.font.size = Pt(14)

								document.add_paragraph(row[column])

					# ----- CVSS -----
					groupColumns = {'cvss(b/v)': 'CVSS (Base Score)', 'cvss(ts/tv)': 'CVSS (Temporal Score)'}
					for key, titleName in groupColumns.items():
						if key in row and len(row[key]) == 2:
							paragraph = document.add_paragraph()
							run = paragraph.add_run('\n' + titleName)
							run.font.color.rgb = scannerColor
							run.font.size = Pt(14)

							document.add_paragraph(row[key][0] + '\t\t' + row[key][1])

					# ----- References -----
					if row['ref']:
						paragraph = document.add_paragraph()
						run = paragraph.add_run('\nReferences')
						run.font.color.rgb = scannerColor
						run.font.size = Pt(14)

						lastRef = ''
						for ref in row['ref']:
							curRef, refID = ref.split(':', 1)
							if curRef != lastRef:

								spacing = '\n'
								if lastRef == '':
									spacing = ''

								paragraph = document.add_paragraph()
								paragraph.style = document.styles['No Spacing']

								paragraph.add_run(spacing + curRef).bold = True
								lastRef = curRef

							paragraph = document.add_paragraph('\t' + refID)
							paragraph.style = document.styles['No Spacing']

					# ----- Links -----
					if row['refURL']:
						paragraph = document.add_paragraph()
						run = paragraph.add_run('\nReferences Links')
						run.font.color.rgb = scannerColor
						run.font.size = Pt(14)

						lastRef = ''
						for refURL in row['refURL']:
							s = refURL.split('(URL)', 1)

							if s[0]:
								ss = s[0].split(':', 1)
								paragraph = document.add_paragraph()
								paragraph.style = document.styles['No Spacing']

								if ss[0] != lastRef:
									lastRef = ss[0]
									paragraph.add_run('{}\n\t{}'.format(ss[0], ss[1])).bold = True
								else:
									paragraph.add_run('\t'+ss[1]).bold = True

								paragraph = document.add_paragraph('\t'+s[1]+'\n')
							else:
								paragraph = document.add_paragraph(s[1])
								lastRef = None

							paragraph.style = document.styles['No Spacing']

					# ----- Others -----
					for newCol in vuln.data:
						if newCol not in ['name', 'summary', 'description', 'solution',
							'port', 'refURL', 'ref', 'extra',
							'severity', 'risk', 'cvss(b/v)', 'cvss(ts/tv)']:

							if row[newCol]:
								paragraph = document.add_paragraph()
								paragraph.style = document.styles['No Spacing']

								run = paragraph.add_run('\n' + newCol.title())
								run.font.color.rgb = scannerColor
								run.font.size = Pt(14)

								if isinstance(row[newCol], list):
									document.add_paragraph(', '.join(filter(None, row[newCol])))
								else:
									document.add_paragraph(row[newCol])

	document.save(fileName)

def compile(colr, outFile, merged):

	if outFile['extension'] == 'docx':
		genDOCX(colr, outFile['filename'], merged)
